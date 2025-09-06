# import tkinter as tk
# from tkinter import filedialog, scrolledtext, messagebox, ttk
# import pyttsx3
# import speech_recognition as sr
# import threading
# import fitz  # PyMuPDF
# from docx import Document
# from deep_translator import GoogleTranslator
# from fpdf import FPDF
# import os
# import nltk
# import string
# import random
# import spacy
# import numpy as np
# from nltk.tokenize import sent_tokenize
# from sklearn.feature_extraction.text import TfidfVectorizer
# from sklearn.metrics.pairwise import cosine_similarity
# from heapq import nlargest
# from spacy.lang.en.stop_words import STOP_WORDS

# nltk.download('punkt')
# nltk.download('wordnet')

# # === Global Context ===
# class ChatContext:
#     def __init__(self):
#         self.text = ""
#         self.sentences = []
#         self.cleaned = []
#         self.vectorizer = None
#         self.tfidf_matrix = None
#         self.chat_history = []
#         self.language = 'english'

#     def load_text(self, text):
#         self.text = text
#         self.sentences = sent_tokenize(text)
#         trans = str.maketrans('', '', string.punctuation)
#         self.cleaned = [s.lower().translate(trans) for s in self.sentences]
#         self.vectorizer = TfidfVectorizer(stop_words='english')
#         self.tfidf_matrix = self.vectorizer.fit_transform(self.cleaned)

# context = ChatContext()

# # === Text-to-Speech ===
# tts_engine = pyttsx3.init()
# tts_engine.setProperty('rate', 160)

# def speak(text, lang='english'):
#     try:
#         if lang != 'english':
#             translated = GoogleTranslator(source='english', target=lang).translate(text)
#         else:
#             translated = text
#         tts_engine.say(translated)
#         tts_engine.runAndWait()
#     except Exception as e:
#         print("TTS Error:", e)

# # === File Handling ===
# def extract_text_from_file(file_path):
#     if file_path.endswith(".pdf"):
#         with fitz.open(file_path) as doc:
#             return "\n".join([page.get_text() for page in doc])
#     elif file_path.endswith(".docx"):
#         doc = Document(file_path)
#         return "\n".join([p.text for p in doc.paragraphs])
#     elif file_path.endswith(".txt"):
#         with open(file_path, "r", encoding="utf-8") as f:
#             return f.read()
#     return ""

# # === Summarizer ===
# def summarize_text(text):
#     nlp = spacy.load("en_core_web_sm")
#     doc = nlp(text)
#     word_freq = {}
#     for word in doc:
#         if word.text.lower() not in STOP_WORDS and word.text not in string.punctuation:
#             word_freq[word.text.lower()] = word_freq.get(word.text.lower(), 0) + 1
#     max_freq = max(word_freq.values())
#     for word in word_freq:
#         word_freq[word] /= max_freq
#     sentence_scores = {}
#     for sent in doc.sents:
#         for word in sent:
#             if word.text.lower() in word_freq:
#                 sentence_scores[sent] = sentence_scores.get(sent, 0) + word_freq[word.text.lower()]
#     summary_len = max(1, int(len(list(doc.sents)) * 0.3))
#     summary = nlargest(summary_len, sentence_scores, key=sentence_scores.get)
#     return ' '.join([s.text for s in summary])

# # === Response Logic ===
# def greeting(sentence):
#     GREETINGS = ("hello", "hi", "greetings", "sup", "what's up", "hey")
#     RESPONSES = ["hi", "hello", "hey there", "*nods*", "Nice to meet you!"]
#     for word in sentence.lower().split():
#         if word in GREETINGS:
#             return random.choice(RESPONSES)
#     return None

# def respond(user_input):
#     if not context.vectorizer or context.tfidf_matrix is None:
#         return "Please load a document first."
#     query = context.vectorizer.transform([user_input])
#     cosine_sim = cosine_similarity(query, context.tfidf_matrix)
#     max_index = np.argmax(cosine_sim)
#     if cosine_sim[0][max_index] == 0:
#         return "Sorry, I couldn't understand that."
#     return context.sentences[max_index]

# # === GUI Setup ===
# class RoboApp:
#     def __init__(self, root):
#         self.root = root
#         self.root.title("ROBO - AI Chatbot")
#         self.root.geometry("750x620")
#         self.root.resizable(True, True)
#         self.setup_gui()

#     def setup_gui(self):
#         self.theme = tk.StringVar(value="light")

#         top_frame = tk.Frame(self.root)
#         top_frame.pack(pady=10)

#         self.language_label = tk.Label(top_frame, text="🌐 Language:")
#         self.language_label.pack(side='left')

#         self.language_dropdown = ttk.Combobox(top_frame, values=["english", "tamil", "hindi", "french", "de", "es"], width=5)
#         self.language_dropdown.set("english")
#         self.language_dropdown.pack(side='left', padx=5)
#         self.language_dropdown.bind("<<ComboboxSelected>>", self.change_language)

#         self.chat_area = scrolledtext.ScrolledText(self.root, wrap=tk.WORD, font=("Segoe UI", 11))
#         self.chat_area.pack(expand=True, fill='both', padx=10, pady=5)
#         self.chat_area.config(state='disabled')

#         input_frame = tk.Frame(self.root)
#         input_frame.pack(fill='x', padx=10, pady=5)

#         self.entry = tk.Entry(input_frame, font=("Segoe UI", 12))
#         self.entry.pack(side='left', fill='x', expand=True, padx=5)
#         self.entry.bind("<Return>", self.send_message)

#         send_btn = tk.Button(input_frame, text="Send", command=self.send_message)
#         send_btn.pack(side='left', padx=5)

#         voice_btn = tk.Button(input_frame, text="🎤", command=self.voice_input)
#         voice_btn.pack(side='left')

#         button_frame = tk.Frame(self.root)
#         button_frame.pack(pady=10)

#         tk.Button(button_frame, text="📂 Load Doc", command=self.load_file).grid(row=0, column=0, padx=5)
#         tk.Button(button_frame, text="📊 Summarize", command=self.summarize_and_export).grid(row=0, column=1, padx=5)
#         tk.Button(button_frame, text="📄 Export Chat", command=self.export_chat).grid(row=0, column=2, padx=5)
#         tk.Button(button_frame, text="🌙 Toggle Theme", command=self.toggle_theme).grid(row=0, column=3, padx=5)

#         self.display_bot("Hello! I'm ROBO. You can type or speak your question.")

#     def change_language(self, event=None):
#         context.language = self.language_dropdown.get()

#     def display_bot(self, text):
#         translated_text = text
#         try:
#             if context.language != 'english':
#                 translated_text = GoogleTranslator(source='english', target=context.language).translate(text)
#         except Exception as e:
#             print("Translation Error:", e)
#         self.chat_area.config(state='normal')
#         self.chat_area.insert(tk.END, f"ROBO: {translated_text}\n")
#         self.chat_area.config(state='disabled')
#         self.chat_area.see(tk.END)
#         threading.Thread(target=speak, args=(text, context.language)).start()

#     def display_user(self, text):
#         self.chat_area.config(state='normal')
#         self.chat_area.insert(tk.END, f"YOU: {text}\n")
#         self.chat_area.config(state='disabled')
#         self.chat_area.see(tk.END)

#     def send_message(self, event=None):
#         user_text = self.entry.get().strip()
#         if not user_text:
#             return
#         self.display_user(user_text)
#         self.entry.delete(0, tk.END)

#         greet = greeting(user_text)
#         response = greet if greet else respond(user_text)
#         context.chat_history.append(("YOU", user_text))
#         context.chat_history.append(("ROBO", response))
#         self.display_bot(response)

#     def voice_input(self):
#         def listen():
#             self.display_bot("🎤 Listening...")
#             r = sr.Recognizer()
#             with sr.Microphone() as source:
#                 try:
#                     audio = r.listen(source, timeout=5)
#                     query = r.recognize_google(audio)
#                     self.entry.delete(0, tk.END)
#                     self.entry.insert(0, query)
#                     self.send_message()
#                 except Exception:
#                     self.display_bot("Sorry, I didn't hear anything.")
#         threading.Thread(target=listen).start()

#     def load_file(self):
#         file_path = filedialog.askopenfilename(filetypes=[("Text", "*.txt"), ("PDF", "*.pdf"), ("Word", "*.docx")])
#         if file_path:
#             text = extract_text_from_file(file_path)
#             context.load_text(text)
#             self.display_bot("Document loaded and ready!")

#     def summarize_and_export(self):
#         if not context.text:
#             self.display_bot("Please load a document first.")
#             return
#         summary = summarize_text(context.text)
#         pdf = FPDF()
#         pdf.add_page()
#         pdf.set_font("Arial", size=12)
#         for line in summary.split('\n'):
#             pdf.multi_cell(0, 10, line)
#         pdf.output("summary.pdf")
#         self.display_bot("Summary saved to summary.pdf.")

#     def export_chat(self):
#         pdf = FPDF()
#         pdf.add_page()
#         pdf.set_font("Arial", size=12)
#         for speaker, msg in context.chat_history:
#             pdf.multi_cell(0, 10, f"{speaker}: {msg}")
#         pdf.output("chat_history.pdf")
#         self.display_bot("Chat exported to chat_history.pdf.")

#     def toggle_theme(self):
#         if self.theme.get() == "light":
#             self.root.tk_setPalette(background="#2E2E2E", foreground="#FFFFFF")
#             self.theme.set("dark")
#         else:
#             self.root.tk_setPalette(background="#F0F0F0", foreground="#000000")
#             self.theme.set("light")

# # === Run App ===
# if __name__ == "__main__":
#     root = tk.Tk()
#     app = RoboApp(root)
#     root.mainloop()
import tkinter as tk
from tkinter import filedialog, scrolledtext, messagebox, ttk
import pyttsx3
import threading
import fitz  # PyMuPDF
from docx import Document
from deep_translator import GoogleTranslator
from fpdf import FPDF
import os
import nltk
import string
import random
import spacy
import numpy as np
import sounddevice as sd
import speech_recognition as sr
from nltk.tokenize import sent_tokenize
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from heapq import nlargest
from spacy.lang.en.stop_words import STOP_WORDS

nltk.download('punkt')
nltk.download('wordnet')

# === Global Context ===
class ChatContext:
    def __init__(self):
        self.text = ""
        self.sentences = []
        self.cleaned = []
        self.vectorizer = None
        self.tfidf_matrix = None
        self.chat_history = []
        self.language = 'english'

    def load_text(self, text):
        self.text = text
        self.sentences = sent_tokenize(text)
        trans = str.maketrans('', '', string.punctuation)
        self.cleaned = [s.lower().translate(trans) for s in self.sentences]
        self.vectorizer = TfidfVectorizer(stop_words='english')
        self.tfidf_matrix = self.vectorizer.fit_transform(self.cleaned)

context = ChatContext()

# === Text-to-Speech ===
tts_engine = pyttsx3.init()
tts_engine.setProperty('rate', 160)

def speak(text, lang='english'):
    try:
        if lang != 'english':
            translated = GoogleTranslator(source='english', target=lang).translate(text)
        else:
            translated = text
        tts_engine.say(translated)
        tts_engine.runAndWait()
    except Exception as e:
        print("TTS Error:", e)

# === File Handling ===
def extract_text_from_file(file_path):
    if file_path.endswith(".pdf"):
        with fitz.open(file_path) as doc:
            return "\n".join([page.get_text() for page in doc])
    elif file_path.endswith(".docx"):
        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    elif file_path.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    return ""

# === Summarizer ===
def summarize_text(text):
    nlp = spacy.load("en_core_web_sm")
    doc = nlp(text)
    word_freq = {}
    for word in doc:
        if word.text.lower() not in STOP_WORDS and word.text not in string.punctuation:
            word_freq[word.text.lower()] = word_freq.get(word.text.lower(), 0) + 1
    max_freq = max(word_freq.values())
    for word in word_freq:
        word_freq[word] /= max_freq
    sentence_scores = {}
    for sent in doc.sents:
        for word in sent:
            if word.text.lower() in word_freq:
                sentence_scores[sent] = sentence_scores.get(sent, 0) + word_freq[word.text.lower()]
    summary_len = max(1, int(len(list(doc.sents)) * 0.3))
    summary = nlargest(summary_len, sentence_scores, key=sentence_scores.get)
    return ' '.join([s.text for s in summary])

# === Response Logic ===
def greeting(sentence):
    GREETINGS = ("hello", "hi", "greetings", "sup", "what's up", "hey")
    RESPONSES = ["hi", "hello", "hey there", "*nods*", "Nice to meet you!"]
    for word in sentence.lower().split():
        if word in GREETINGS:
            return random.choice(RESPONSES)
    return None

def respond(user_input):
    if not context.vectorizer or context.tfidf_matrix is None:
        return "Please load a document first."
    query = context.vectorizer.transform([user_input])
    cosine_sim = cosine_similarity(query, context.tfidf_matrix)
    max_index = np.argmax(cosine_sim)
    if cosine_sim[0][max_index] == 0:
        return "Sorry, I couldn't understand that."
    return context.sentences[max_index]

# === GUI Setup ===
class RoboApp:
    def __init__(self, root):
        self.root = root
        self.root.title("ROBO - AI Chatbot")
        self.root.geometry("750x620")
        self.root.resizable(True, True)
        self.setup_gui()

    def setup_gui(self):
        self.theme = tk.StringVar(value="light")

        top_frame = tk.Frame(self.root)
        top_frame.pack(pady=10)

        self.language_label = tk.Label(top_frame, text="🌐 Language:")
        self.language_label.pack(side='left')

        self.language_dropdown = ttk.Combobox(top_frame, values=["english", "tamil", "hindi", "french", "de", "es"], width=5)
        self.language_dropdown.set("english")
        self.language_dropdown.pack(side='left', padx=5)
        self.language_dropdown.bind("<<ComboboxSelected>>", self.change_language)

        self.chat_area = scrolledtext.ScrolledText(self.root, wrap=tk.WORD, font=("Segoe UI", 11))
        self.chat_area.pack(expand=True, fill='both', padx=10, pady=5)
        self.chat_area.config(state='disabled')

        input_frame = tk.Frame(self.root)
        input_frame.pack(fill='x', padx=10, pady=5)

        self.entry = tk.Entry(input_frame, font=("Segoe UI", 12))
        self.entry.pack(side='left', fill='x', expand=True, padx=5)
        self.entry.bind("<Return>", self.send_message)

        send_btn = tk.Button(input_frame, text="Send", command=self.send_message)
        send_btn.pack(side='left', padx=5)

        voice_btn = tk.Button(input_frame, text="🎤", command=self.voice_input)
        voice_btn.pack(side='left')

        button_frame = tk.Frame(self.root)
        button_frame.pack(pady=10)

        tk.Button(button_frame, text="📂 Load Doc", command=self.load_file).grid(row=0, column=0, padx=5)
        tk.Button(button_frame, text="📊 Summarize", command=self.summarize_and_export).grid(row=0, column=1, padx=5)
        tk.Button(button_frame, text="📄 Export Chat", command=self.export_chat).grid(row=0, column=2, padx=5)
        tk.Button(button_frame, text="🌙 Toggle Theme", command=self.toggle_theme).grid(row=0, column=3, padx=5)

        self.display_bot("Hello! I'm ROBO. You can type or speak your question.")

    def change_language(self, event=None):
        context.language = self.language_dropdown.get()

    def display_bot(self, text):
        translated_text = text
        try:
            if context.language != 'english':
                translated_text = GoogleTranslator(source='english', target=context.language).translate(text)
        except Exception as e:
            print("Translation Error:", e)
        self.chat_area.config(state='normal')
        self.chat_area.insert(tk.END, f"ROBO: {translated_text}\n")
        self.chat_area.config(state='disabled')
        self.chat_area.see(tk.END)
        threading.Thread(target=speak, args=(text, context.language)).start()

    def display_user(self, text):
        self.chat_area.config(state='normal')
        self.chat_area.insert(tk.END, f"YOU: {text}\n")
        self.chat_area.config(state='disabled')
        self.chat_area.see(tk.END)

    def send_message(self, event=None):
        user_text = self.entry.get().strip()
        if not user_text:
            return
        self.display_user(user_text)
        self.entry.delete(0, tk.END)

        greet = greeting(user_text)
        response = greet if greet else respond(user_text)
        context.chat_history.append(("YOU", user_text))
        context.chat_history.append(("ROBO", response))
        self.display_bot(response)

    def voice_input(self):
        def listen():
            self.display_bot("🎤 Listening... Please speak for 5 seconds.")
            fs = 16000
            duration = 5
            try:
                recording = sd.rec(int(duration * fs), samplerate=fs, channels=1, dtype='int16')
                sd.wait()
                audio_data = sr.AudioData(recording.tobytes(), fs, 2)
                recognizer = sr.Recognizer()
                query = recognizer.recognize_google(audio_data)
                self.entry.delete(0, tk.END)
                self.entry.insert(0, query)
                self.send_message()
            except Exception as e:
                self.display_bot("Sorry, I couldn't capture your voice.")
                print("Voice Input Error:", e)
        threading.Thread(target=listen).start()

    def load_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Text", "*.txt"), ("PDF", "*.pdf"), ("Word", "*.docx")])
        if file_path:
            text = extract_text_from_file(file_path)
            context.load_text(text)
            self.display_bot("Document loaded and ready!")

    def summarize_and_export(self):
        if not context.text:
            self.display_bot("Please load a document first.")
            return
        summary = summarize_text(context.text)
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for line in summary.split('\n'):
            pdf.multi_cell(0, 10, line)
        pdf.output("summary.pdf")
        self.display_bot("Summary saved to summary.pdf.")

    def export_chat(self):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for speaker, msg in context.chat_history:
            pdf.multi_cell(0, 10, f"{speaker}: {msg}")
        pdf.output("chat_history.pdf")
        self.display_bot("Chat exported to chat_history.pdf.")

    def toggle_theme(self):
        if self.theme.get() == "light":
            self.root.tk_setPalette(background="#2E2E2E", foreground="#FFFFFF")
            self.theme.set("dark")
        else:
            self.root.tk_setPalette(background="#F0F0F0", foreground="#000000")
            self.theme.set("light")

# === Run App ===
if __name__ == "__main__":
    root = tk.Tk()
    app = RoboApp(root)
    root.mainloop()
